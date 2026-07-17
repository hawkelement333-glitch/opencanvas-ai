from __future__ import annotations

import io
import re
from dataclasses import dataclass, replace

from docx import Document as DocxDocument
from pypdf import PdfReader

from opencanvas_api.core.config import Settings
from opencanvas_api.services.documents.errors import DocumentExtractionError

_MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
_MULTIPLE_BLANK_LINES = re.compile(r"\n{3,}")
_INLINE_WHITESPACE = re.compile(r"[\t\v\f ]+")


@dataclass(frozen=True, slots=True)
class ExtractedSegment:
    text: str
    page_number: int | None = None
    heading: str | None = None
    char_start: int = 0
    char_end: int = 0


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    text: str
    segments: tuple[ExtractedSegment, ...]
    page_count: int | None


def extract_document(*, file_type: str, content: bytes, settings: Settings) -> ExtractedDocument:
    try:
        if file_type == "pdf":
            segments, page_count = _extract_pdf(content, settings)
        elif file_type == "txt":
            segments, page_count = [_text_segment(content)], None
        elif file_type == "markdown":
            segments, page_count = _extract_markdown(content), None
        elif file_type == "docx":
            segments, page_count = _extract_docx(content), None
        else:
            raise DocumentExtractionError("This document type cannot be extracted.")
    except DocumentExtractionError:
        raise
    except Exception as exc:
        raise DocumentExtractionError("The document is malformed or could not be read.") from exc
    return _assemble(segments, page_count=page_count, settings=settings)


def normalize_extracted_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n").replace("\u00a0", " ")
    normalized_lines = [_INLINE_WHITESPACE.sub(" ", line).rstrip() for line in value.split("\n")]
    return _MULTIPLE_BLANK_LINES.sub("\n\n", "\n".join(normalized_lines)).strip()


def _extract_pdf(content: bytes, settings: Settings) -> tuple[list[ExtractedSegment], int]:
    try:
        reader = PdfReader(io.BytesIO(content), strict=False)
    except Exception as exc:
        raise DocumentExtractionError("The PDF is malformed and could not be opened.") from exc
    page_count = len(reader.pages)
    if page_count == 0:
        raise DocumentExtractionError("The PDF does not contain any pages.")
    if page_count > settings.document_max_pdf_pages:
        raise DocumentExtractionError("The PDF contains too many pages to process safely.")

    segments: list[ExtractedSegment] = []
    extracted_characters = 0
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            raise DocumentExtractionError(
                f"Text could not be extracted from PDF page {page_number}."
            ) from exc
        extracted_characters += len(text)
        if extracted_characters > settings.document_max_extracted_characters:
            raise DocumentExtractionError(
                "The PDF contains too much extracted text to process safely."
            )
        if normalized := normalize_extracted_text(text):
            segments.append(ExtractedSegment(text=normalized, page_number=page_number))
    return segments, page_count


def _text_segment(content: bytes) -> ExtractedSegment:
    return ExtractedSegment(text=content.decode("utf-8-sig"))


def _extract_markdown(content: bytes) -> list[ExtractedSegment]:
    text = content.decode("utf-8-sig")
    segments: list[ExtractedSegment] = []
    heading: str | None = None
    section_lines: list[str] = []

    def finish_section() -> None:
        normalized = normalize_extracted_text("\n".join(section_lines))
        if normalized or heading:
            prefix = f"{heading}\n\n" if heading else ""
            segments.append(ExtractedSegment(text=f"{prefix}{normalized}".strip(), heading=heading))

    for line in text.splitlines():
        match = _MARKDOWN_HEADING.match(line.strip())
        if match:
            finish_section()
            heading = match.group(2).strip()[:500]
            section_lines = []
        else:
            section_lines.append(line)
    finish_section()
    if not segments and heading:
        segments.append(ExtractedSegment(text=heading, heading=heading))
    return segments


def _extract_docx(content: bytes) -> list[ExtractedSegment]:
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError("The DOCX document could not be opened.") from exc

    segments: list[ExtractedSegment] = []
    heading: str | None = None
    section_lines: list[str] = []

    def finish_section() -> None:
        normalized = normalize_extracted_text("\n\n".join(section_lines))
        if normalized or heading:
            prefix = f"{heading}\n\n" if heading else ""
            segments.append(ExtractedSegment(text=f"{prefix}{normalized}".strip(), heading=heading))

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.lower().startswith("heading"):
            finish_section()
            heading = text[:500]
            section_lines = []
        else:
            section_lines.append(text)
    for table in document.tables:
        for row in table.rows:
            values = [normalize_extracted_text(cell.text) for cell in row.cells]
            if any(values):
                section_lines.append(" | ".join(values))
    finish_section()
    if not segments and heading:
        segments.append(ExtractedSegment(text=heading, heading=heading))
    return segments


def _assemble(
    segments: list[ExtractedSegment], *, page_count: int | None, settings: Settings
) -> ExtractedDocument:
    normalized_segments: list[ExtractedSegment] = []
    pieces: list[str] = []
    cursor = 0
    for segment in segments:
        normalized = normalize_extracted_text(segment.text)
        if not normalized:
            continue
        if pieces:
            pieces.append("\n\n")
            cursor += 2
        start = cursor
        pieces.append(normalized)
        cursor += len(normalized)
        normalized_segments.append(
            replace(segment, text=normalized, char_start=start, char_end=cursor)
        )
        if cursor > settings.document_max_extracted_characters:
            raise DocumentExtractionError("The document contains too much text to process safely.")
    full_text = "".join(pieces)
    if not full_text.strip():
        raise DocumentExtractionError(
            "No readable text was found. Image-only documents require OCR, which is not enabled."
        )
    return ExtractedDocument(
        text=full_text,
        segments=tuple(normalized_segments),
        page_count=page_count,
    )
