from __future__ import annotations

import io

from docx import Document as DocxDocument
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject


def make_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Launch brief", level=1)
    document.add_paragraph("The launch codename is Aurora Finch. The pilot began on 14 April 2026.")
    document.add_heading("Scope", level=2)
    document.add_paragraph("The pilot includes twelve research teams in Chicago.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def make_pdf_bytes(
    text: str | list[str] = "The launch codename is Aurora Finch.",
) -> bytes:
    writer = PdfWriter()
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    page_texts = [text] if isinstance(text, str) else text
    for page_text in page_texts:
        page = writer.add_blank_page(width=612, height=792)
        resources = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_reference})}
        )
        stream = DecodedStreamObject()
        escaped = page_text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1"))
        page[NameObject("/Resources")] = resources
        page[NameObject("/Contents")] = writer._add_object(stream)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()
